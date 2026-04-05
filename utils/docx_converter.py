"""
One-time utility: convert a .docx resume to master_resume.md

Usage:
    python utils/docx_converter.py path/to/resume.docx
    python utils/docx_converter.py path/to/resume.docx --output data/master_resume.md
"""

import sys
import argparse
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def docx_to_markdown(docx_path: str, output_path: str = "data/master_resume.md") -> str:
    doc = Document(docx_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        style_name = para.style.name.lower()

        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif para.style.name.startswith("List"):
            lines.append(f"- {text}")
        else:
            lines.append(text)

    # Also extract tables (for skills grids etc.)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
        lines.append("")

    markdown = "\n".join(lines)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(markdown, encoding="utf-8")
    print(f"✓ Resume saved to {output_path} ({len(doc.paragraphs)} paragraphs)")
    return markdown


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert .docx resume to markdown")
    parser.add_argument("docx_path", help="Path to your .docx resume file")
    parser.add_argument("--output", default="data/master_resume.md", help="Output path for markdown file")
    args = parser.parse_args()

    if not Path(args.docx_path).exists():
        print(f"ERROR: File not found: {args.docx_path}")
        sys.exit(1)

    docx_to_markdown(args.docx_path, args.output)
